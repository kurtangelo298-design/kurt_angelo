from flask import Flask, render_template_string, request, jsonify, make_response
import psycopg2
import os
import datetime
import barcode
from barcode.writer import ImageWriter
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# ===================== DATABASE — NEON POSTGRESQL =====================
DB_URI = os.environ.get("DATABASE_URL", "")

# ===================== LOGIN CREDENTIALS =====================
USERNAME = "slsu"
PASSWORD = "jge"

# ===================== DATABASE INIT =====================
def init_db():
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        full_name TEXT NOT NULL,
        department TEXT NOT NULL,
        contact_number TEXT,
        year_level TEXT NOT NULL,
        student_number TEXT NOT NULL UNIQUE,
        registered_at TEXT NOT NULL
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS attendance (
        id SERIAL PRIMARY KEY,
        user_id INTEGER NOT NULL REFERENCES users(id),
        time_in TEXT,
        time_out TEXT,
        scan_date TEXT NOT NULL
    )""")
    conn.commit()
    conn.close()
    print("✅ Database Ready — PERMANENT! HINDI NA MABUBURA!")

init_db()

DEPARTMENTS = ["CT", "FBT", "BSED", "BEED", "BSFI", "BSBA", "EMPLOYEE"]
YEAR_LEVELS = ["1st Year", "2nd Year", "3rd Year", "4th Year", "5th Year", "N/A"]

def generate_barcode_img(student_number):
    code128 = barcode.get_barcode_class('code128')
    barcode_inst = code128(student_number, writer=ImageWriter())
    buffer = BytesIO()
    barcode_inst.write(buffer, options={"write_text": True, "module_width": 0.3, "module_height": 8})
    buffer.seek(0)
    img_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    return img_base64

def is_logged_in():
    return request.cookies.get('logged_in') == 'true'

# ===================== LOGIN =====================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        uname = request.form.get('username', '').strip()
        pword = request.form.get('password', '').strip()
        if uname == USERNAME and pword == PASSWORD:
            resp = make_response('<script>window.location="/";</script>')
            resp.set_cookie('logged_in', 'true', max_age=31536000)
            return resp
        return render_template_string("""
            <html><body style="background:#f0f4f8; font-family:Arial; display:flex; justify-content:center; align-items:center; height:100vh; margin:0;">
                <div style="background:white; padding:30px; border-radius:12px; box-shadow:0 4px 12px rgba(0,0,0,0.1); width:320px;">
                    <h2 style="text-align:center; color:#2c3e50; margin-bottom:20px;">🔐 Admin Login</h2>
                    <p style="color:red; text-align:center; margin-bottom:15px;">❌ Wrong username or password!</p>
                    <form method="POST">
                        <div style="margin-bottom:15px;">
                            <label style="display:block; margin-bottom:5px; font-weight:bold;">Username</label>
                            <input type="text" name="username" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                        </div>
                        <div style="margin-bottom:20px;">
                            <label style="display:block; margin-bottom:5px; font-weight:bold;">Password</label>
                            <input type="password" name="password" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                        </div>
                        <button type="submit" style="width:100%; padding:12px; background:#28a745; color:white; border:none; border-radius:6px; font-size:16px; cursor:pointer;">Login</button>
                    </form>
                </div>
            </body></html>
        """)
    return render_template_string("""
        <html><body style="background:#f0f4f8; font-family:Arial; display:flex; justify-content:center; align-items:center; height:100vh; margin:0;">
            <div style="background:white; padding:30px; border-radius:12px; box-shadow:0 4px 12px rgba(0,0,0,0.1); width:320px;">
                <h2 style="text-align:center; color:#2c3e50; margin-bottom:20px;">🔐 Admin Login</h2>
                <form method="POST">
                    <div style="margin-bottom:15px;">
                        <label style="display:block; margin-bottom:5px; font-weight:bold;">Username</label>
                        <input type="text" name="username" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                    </div>
                    <div style="margin-bottom:20px;">
                        <label style="display:block; margin-bottom:5px; font-weight:bold;">Password</label>
                        <input type="password" name="password" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                    </div>
                    <button type="submit" style="width:100%; padding:12px; background:#28a745; color:white; border:none; border-radius:6px; font-size:16px; cursor:pointer;">Login</button>
                </form>
            </div>
        </body></html>
    """)

@app.route('/logout')
def logout():
    resp = make_response('<script>window.location="/login";</script>')
    resp.set_cookie('logged_in', '', expires=0)
    return resp

# ===================== MAIN DASHBOARD =====================
@app.route('/')
def index():
    if not is_logged_in():
        return '<script>window.location="/login";</script>'
    return render_template_string("""
<!DOCTYPE html>
<html>
<head>
    <title>📚 Library Attendance System</title>
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <style>
        *{box-sizing:border-box; margin:0; padding:0; font-family:'Segoe UI', Arial, sans-serif;}
        body{background:#f0f4f8; padding:20px;}
        .container{max-width:1200px; margin:0 auto;}
        h1{text-align:center; color:#2c3e50; margin-bottom:25px;}
        .logout-btn{position:absolute; top:20px; right:20px; background:#dc3545; color:white; border:none; padding:10px 18px; border-radius:6px; cursor:pointer; font-size:15px;}
        .tabs{display:flex; flex-wrap:wrap; gap:10px; margin-bottom:25px;}
        .tab{padding:12px 20px; background:white; border:none; border-radius:8px; cursor:pointer; font-size:15px; font-weight:600; color:#555; transition:all 0.2s;}
        .tab:hover{background:#e9ecef;}
        .tab.active{background:#28a745; color:white;}
        .tab-content{background:white; padding:25px; border-radius:12px; box-shadow:0 2px 10px rgba(0,0,0,0.08); display:none;}
        .tab-content.active{display:block;}
        h2{color:#2c3e50; margin-bottom:20px; border-bottom:2px solid #eee; padding-bottom:10px;}
        input, select{width:100%; padding:12px; margin:8px 0 15px 0; border:1px solid #ddd; border-radius:6px; font-size:15px;}
        button{padding:12px 24px; background:#28a745; color:white; border:none; border-radius:6px; font-size:15px; cursor:pointer; margin:5px 5px 5px 0;}
        button:hover{background:#218838;}
        .btn-red{background:#dc3545;}
        .btn-red:hover{background:#c82333;}
        .btn-blue{background:#007bff;}
        .btn-blue:hover{background:#0056b3;}
        table{width:100%; border-collapse:collapse; margin-top:15px;}
        th, td{padding:12px; text-align:left; border-bottom:1px solid #eee;}
        th{background:#f8f9fa; font-weight:bold; color:#2c3e50;}
        .scan-box{text-align:center; padding:30px; background:#f8f9fa; border-radius:10px; margin-bottom:20px;}
        .scan-input{font-size:22px; text-align:center; padding:15px;}
        .success{color:green; font-weight:bold; margin-top:15px;}
        .error{color:red; font-weight:bold; margin-top:15px;}
        .barcode-img{max-width:250px; margin:15px 0;}
        .edit-form{background:#f8f9fa; padding:20px; border-radius:8px; margin-top:20px; display:none;}
    </style>
</head>
<body>
    <div class="container">
        <button class="logout-btn" onclick="window.location='/logout'">🚪 Logout</button>
        <h1>📚 Library Attendance System</h1>

        <div class="tabs">
            <button class="tab active" onclick="showTab('scan')">📱 Scan / Attendance</button>
            <button class="tab" onclick="showTab('register')">👤 Register Student</button>
            <button class="tab" onclick="showTab('students')">📋 All Students</button>
            <button class="tab" onclick="showTab('records')">📊 Attendance Records</button>
        </div>

        <!-- SCAN TAB -->
        <div id="scan" class="tab-content active">
            <h2>📱 Scan Student Barcode</h2>
            <div class="scan-box">
                <input type="text" id="scanInput" class="scan-input" placeholder="Scan barcode or type student number..." autofocus>
                <div id="scanResult"></div>
            </div>
        </div>

        <!-- REGISTER TAB -->
        <div id="register" class="tab-content">
            <h2>👤 Register New Student</h2>
            <form id="registerForm">
                <label><b>Full Name:</b></label>
                <input type="text" name="full_name" required placeholder="Juan Dela Cruz">
                
                <label><b>Department:</b></label>
                <select name="department" required>
                    <option value="">-- Select Department --</option>
                    {% for d in depts %}<option value="{{d}}">{{d}}</option>{% endfor %}
                </select>
                
                <label><b>Year Level:</b></label>
                <select name="year_level" required>
                    <option value="">-- Select Year --</option>
                    {% for y in years %}<option value="{{y}}">{{y}}</option>{% endfor %}
                </select>
                
                <label><b>Student Number:</b></label>
                <input type="text" name="student_number" required placeholder="e.g. 2024-0001">
                
                <label><b>Contact Number:</b></label>
                <input type="text" name="contact_number" placeholder="09123456789">
                
                <button type="submit">✅ Register & Generate Barcode</button>
            </form>
            <div id="registerResult"></div>
        </div>

        <!-- STUDENTS TAB -->
        <div id="students" class="tab-content">
            <h2>📋 Registered Students</h2>
            <button class="btn-blue" onclick="loadStudents()">🔄 Refresh List</button>
            <div id="studentsList"></div>
        </div>

        <!-- RECORDS TAB -->
        <div id="records" class="tab-content">
            <h2>📊 Attendance Records — Today</h2>
            <button class="btn-blue" onclick="loadRecords()">🔄 Refresh Records</button>
            <button class="btn-blue" onclick="exportRecords()">📄 Export to Word</button>
            <div id="recordsList"></div>
        </div>
    </div>

<script>
let currentEditId = null;

function showTab(tabName){
    document.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));
    document.querySelectorAll('.tab-content').forEach(c=>c.classList.remove('active'));
    event.target.classList.add('active');
    document.getElementById(tabName).classList.add('active');
    if(tabName==='students') loadStudents();
    if(tabName==='records') loadRecords();
    if(tabName==='scan') setTimeout(()=>document.getElementById('scanInput').focus(),100);
}

// SCAN / ATTENDANCE
document.getElementById('scanInput').addEventListener('keypress', function(e){
    if(e.key==='Enter') submitScan();
});

function submitScan(){
    const code = document.getElementById('scanInput').value.trim();
    if(!code) return;
    fetch('/scan', {
        method:'POST',
        headers:{'Content-Type':'application/x-www-form-urlencoded'},
        body:'student_number='+encodeURIComponent(code)
    }).then(r=>r.json()).then(d=>{
        const res = document.getElementById('scanResult');
        if(d.success){
            res.innerHTML = <div class="success">✅ ${d.message}</div>;
        }else{
            res.innerHTML = <div class="error">❌ ${d.message}</div>;
        }
        document.getElementById('scanInput').value='';
        setTimeout(()=>document.getElementById('scanInput').focus(),100);
    });
}

// REGISTER
document.getElementById('registerForm').addEventListener('submit', function(e){
    e.preventDefault();
    const fd = new FormData(this);
    fetch('/register', {method:'POST', body:fd})
    .then(r=>r.json()).then(d=>{
        const res = document.getElementById('registerResult');
        if(d.success){
            res.innerHTML = `<div class="success"><h3>✅ Registered Successfully!</h3>
                <p><b>Name:</b> ${d.student.full_name}</p>
                <p><b>Student No:</b> ${d.student.student_number}</p>
                <p><b>Barcode:</b></p>
                <img src="data:image/png;base64,${d.barcode}" class="barcode-img"><br>
                <button onclick="window.open('/print-barcode/${d.student.id}')">🖨️ Print Barcode</button>
                </div>`;
            this.reset();
        }else{
            res.innerHTML = <div class="error">❌ ${d.message}</div>;
        }
    });
});

// LOAD STUDENTS
function loadStudents(){
    fetch('/students').then(r=>r.text()).then(h=>{
        document.getElementById('studentsList').innerHTML = h;
    });
}

// LOAD RECORDS
function loadRecords(){
    fetch('/records').then(r=>r.text()).then(h=>{
        document.getElementById('recordsList').innerHTML = h;
    });
}

// EDIT STUDENT
function showEditForm(id, name, dept, year, num, contact){
    currentEditId = id;
    document.getElementById('edit_id').value = id;
    document.getElementById('edit_full_name').value = name;
    document.getElementById('edit_department').value = dept;
    document.getElementById('edit_year_level').value = year;
    document.getElementById('edit_student_number').value = num;
    document.getElementById('edit_contact_number').value = contact;
    document.getElementById('editForm').style.display = 'block';
}

function hideEditForm(){
    document.getElementById('editForm').style.display = 'none';
    currentEditId = null;
}

// EXPORT
function exportRecords(){
    window.location.href = '/export-word';
}
</script>

<!-- EDIT FORM -->
<div id="editForm" class="edit-form" style="display:none;">
    <h3>✏️ Edit Student Info</h3>
    <form id="editStudentForm">
        <input type="hidden" id="edit_id">
        <label><b>Full Name:</b></label>
        <input type="text" id="edit_full_name" required>
        <label><b>Department:</b></label>
        <select id="edit_department" required>
            <option value="CT">CT</option>
            <option value="FBT">FBT</option>
            <option value="BSED">BSED</option>
            <option value="BEED">BEED</option>
            <option value="BSFI">BSFI</option>
            <option value="BSBA">BSBA</option>
            <option value="EMPLOYEE">EMPLOYEE</option>
        </select>
        <label><b>Year Level:</b></label>
        <select id="edit_year_level" required>
            <option value="1st Year">1st Year</option>
            <option value="2nd Year">2nd Year</option>
            <option value="3rd Year">3rd Year</option>
            <option value="4th Year">4th Year</option>
            <option value="5th Year">5th Year</option>
            <option value="N/A">N/A</option>
        </select>
        <label><b>Student Number:</b></label>
        <input type="text" id="edit_student_number" required>
        <label><b>Contact Number:</b></label>
        <input type="text" id="edit_contact_number">
        <button type="button" onclick="saveEdit()">💾 Save Changes</button>
        <button type="button" class="btn-red" onclick="hideEditForm()">❌ Cancel</button>
    </form>
</div>

<script>
function saveEdit(){
    const id = document.getElementById('edit_id').value;
    const data = {
        id: id,
        full_name: document.getElementById('edit_full_name').value,
        department: document.getElementById('edit_department').value,
        year_level: document.getElementById('edit_year_level').value,
        student_number: document.getElementById('edit_student_number').value,
        contact_number: document.getElementById('edit_contact_number').value
    };
    fetch('/update-student', {
        method:'POST',
        headers:{'Content-Type':'application/json'},
        body: JSON.stringify(data)
    }).then(r=>r.json()).then(d=>{
        if(d.success){
            alert('✅ Updated successfully!');
            hideEditForm();
            loadStudents();
        }else{
            alert('❌ Error: ' + d.message);
        }
    });
}
</script>

</body>
</html>
    """, depts=DEPARTMENTS, years=YEAR_LEVELS)

# ===================== SCAN ENDPOINT =====================
@app.route('/scan', methods=['POST'])
def scan():
    if not is_logged_in():
        return jsonify({"success": False, "message": "Please login first"})
    
    student_number = request.form.get('student_number', '').strip().upper()
    now = datetime.datetime.now()
    today = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%I:%M %p")  # Oras lang, walang date
    
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    
    # Hanapin ang student
    c.execute("SELECT id, full_name FROM users WHERE UPPER(student_number) = %s", (student_number,))
    user = c.fetchone()
    if not user:
        conn.close()
        return jsonify({"success": False, "message": f"Student not found: {student_number}"})
    
    user_id, full_name = user
    
    # Check kung may active IN na walang OUT
    c.execute("SELECT id FROM attendance WHERE user_id = %s AND scan_date = %s AND time_out IS NULL", (user_id, today))
    active = c.fetchone()
    
    if active:
        # TIME OUT
        c.execute("UPDATE attendance SET time_out = %s WHERE id = %s", (time_str, active[0]))
        msg = f"⏰ TIME OUT — {full_name} — {time_str}"
    else:
        # TIME IN
        c.execute("INSERT INTO attendance (user_id, time_in, scan_date) VALUES (%s, %s, %s)", (user_id, time_str, today))
        msg = f"✅ TIME IN — {full_name} — {time_str}"
    
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": msg})

# ===================== REGISTER ENDPOINT =====================
@app.route('/register', methods=['POST'])
def register():
    if not is_logged_in():
        return jsonify({"success": False, "message": "Please login first"})
    
    full_name = request.form.get('full_name', '').strip()
    department = request.form.get('department', '').strip()
    year_level = request.form.get('year_level', '').strip()
    student_number = request.form.get('student_number', '').strip().upper()
    contact_number = request.form.get('contact_number', '').strip()
    registered_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    if not all([full_name, department, year_level, student_number]):
        return jsonify({"success": False, "message": "Fill all required fields"})
    
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    try:
        c.execute("""INSERT INTO users 
            (full_name, department, contact_number, year_level, student_number, registered_at)
            VALUES (%s, %s, %s, %s, %s, %s) RETURNING id""",
            (full_name, department, contact_number, year_level, student_number, registered_at))
        user_id = c.fetchone()[0]
        conn.commit()
        
        barcode_b64 = generate_barcode_img(student_number)
        
        return jsonify({
            "success": True,
            "student": {
                "id": user_id,
                "full_name": full_name,
                "student_number": student_number
            },
            "barcode": barcode_b64
        })
    except psycopg2.IntegrityError:
        conn.rollback()
        return jsonify({"success": False, "message": "Student Number already exists!"})
    finally:
        conn.close()

# ===================== STUDENTS LIST =====================
@app.route('/students')
def students_list():
    if not is_logged_in():
        return "Unauthorized"
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("SELECT id, full_name, department, year_level, student_number, contact_number FROM users ORDER BY full_name")
    students = c.fetchall()
    conn.close()
    
    html = "<table><tr><th>Name</th><th>Dept</th><th>Year</th><th>Student No.</th><th>Contact</th><th>Action</th></tr>"
    for s in students:
        html += f"""<tr>
            <td>{s[1]}</td><td>{s[2]}</td><td>{s[3]}</td><td>{s[4]}</td><td>{s[5] or '-'}</td>
            <td><button onclick="showEditForm({s[0]}, '{s[1]}', '{s[2]}', '{s[3]}', '{s[4]}', '{s[5] or ''}')">✏️ Edit</button>
            <a href="/print-barcode/{s[0]}" target="_blank"><button>🖨️ Barcode</button></a></td>
        </tr>"""
    html += "</table>"
    return html

# ===================== UPDATE STUDENT =====================
@app.route('/update-student', methods=['POST'])
def update_student():
    if not is_logged_in():
        return jsonify({"success": False})
    data = request.get_json()
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    try:
        c.execute("""UPDATE users SET full_name=%s, department=%s, year_level=%s, 
            student_number=%s, contact_number=%s WHERE id=%s""",
            (data['full_name'], data['department'], data['year_level'], 
             data['student_number'].upper(), data['contact_number'], data['id']))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": str(e)})
    finally:
        conn.close()

# ===================== ATTENDANCE RECORDS =====================
@app.route('/records')
def records():
    if not is_logged_in():
        return "Unauthorized"
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""SELECT u.full_name, u.department, a.time_in, a.time_out 
        FROM attendance a JOIN users u ON a.user_id = u.id 
        WHERE a.scan_date = %s ORDER BY a.id DESC""", (today,))
    recs = c.fetchall()
    conn.close()
    
    html = f"<h3>Attendance for {today}</h3><table><tr><th>Name</th><th>Dept</th><th>Time In</th><th>Time Out</th></tr>"
    for r in recs:
        html += f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[3] or '-'}</td></tr>"
    html += "</table>"
    return html

# ===================== EXPORT TO WORD =====================
@app.route('/export-word')
def export_word():
    if not is_logged_in():
        return "Unauthorized"
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""SELECT u.full_name, u.department, a.time_in, a.time_out 
        FROM attendance a JOIN users u ON a.user_id = u.id 
        WHERE a.scan_date = %s ORDER BY a.id DESC""", (today,))
    recs = c.fetchall()
    conn.close()
    
    doc = Document()
    doc.add_heading(f'Library Attendance — {today}', 0)
    doc.add_paragraph(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Name'
    hdr[1].text = 'Department'
    hdr[2].text = 'Time In'
    hdr[3].text = 'Time Out'
    
    for r in recs:
        row = table.add_row().cells
        row[0].text = r[0]
        row[1].text = r[1]
        row[2].text = r[2]
        row[3].text = r[3] or '-'
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    resp = make_response(buffer.read())
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    resp.headers['Content-Disposition'] = f'attachment; filename=attendance_{today}.docx'
    return resp

# ===================== PRINT BARCODE =====================
@app.route('/print-barcode/<int:user_id>')
def print_barcode(user_id):
    if not is_logged_in():
        return "Unauthorized"
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("SELECT full_name, student_number FROM users WHERE id = %s", (user_id,))
    user = c.fetchone()
    conn.close()
    if not user:
        return "User not found"
    
    barcode_b64 = generate_barcode_img(user[1])
    
    return f"""
    <html>
    <head><title>Barcode — {user[1]}</title>
    <style>
        body {{ text-align: center; padding: 40px; font-family: Arial; }}
        .barcode {{ max-width: 250px; margin: 20px auto; }}
        .info {{ font-size: 18px; margin: 10px 0; }}
        @media print {{ button {{ display: none; }} }}
    </style>
    </head>
    <body>
        <h2>{user[0]}</h2>
        <div class="info">Student No: {user[1]}</div>
        <img src="data:image/png;base64,{barcode_b64}" class="barcode">
        <br><br>
        <button onclick="window.print()" style="padding:12px 24px; font-size:16px; cursor:pointer;">🖨️ Print</button>
    </body>
    </html>
    """

if _name_ == '_main_':
    app.run(host='0.0.0.0', port=5000, debug=False)
from flask import Flask, render_template_string, request, jsonify, make_response
import psycopg2
import os
import datetime
import barcode
from barcode.writer import ImageWriter
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(_name_)

# ===================== DATABASE — NEON POSTGRESQL =====================
DB_URI = os.environ.get("DATABASE_URL", "")

# ===================== LOGIN CREDENTIALS =====================
USERNAME = "slsu"
PASSWORD = "jge"

# ===================== DATABASE INIT =====================
def init_db():
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        full_name TEXT NOT NULL,
        department TEXT NOT NULL,
        contact_number TEXT,
        year_level TEXT NOT NULL,
        student_number TEXT NOT NULL UNIQUE,
        registered_at TEXT NOT NULL
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS attendance (
        id SERIAL PRIMARY KEY,
        user_id INTEGER NOT NULL REFERENCES users(id),
        time_in TEXT,
        time_out TEXT,
        scan_date TEXT NOT NULL
    )""")
    conn.commit()
    conn.close()
    print("✅ Database Ready — PERMANENT! HINDI NA MABUBURA!")

init_db()

DEPARTMENTS = ["CT", "FBT", "BSED", "BEED", "BSFI", "BSBA", "EMPLOYEE"]
YEAR_LEVELS = ["1st Year", "2nd Year", "3rd Year", "4th Year", "5th Year", "N/A"]

def generate_barcode_img(student_number):
    code128 = barcode.get_barcode_class('code128')
    barcode_inst = code128(student_number, writer=ImageWriter())
    buffer = BytesIO()
    barcode_inst.write(buffer, options={"write_text": True, "module_width": 0.3, "module_height": 8})
    buffer.seek(0)
    img_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    return img_base64

def is_logged_in():
    return request.cookies.get('logged_in') == 'true'

# ===================== LOGIN =====================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        uname = request.form.get('username', '').strip()
        pword = request.form.get('password', '').strip()
        if uname == USERNAME and pword == PASSWORD:
            resp = make_response('<script>window.location="/";</script>')
            resp.set_cookie('logged_in', 'true', max_age=31536000)
            return resp
        return render_template_string("""
            <html><body style="background:#f0f4f8; font-family:Arial; display:flex; justify-content:center; align-items:center; height:100vh; margin:0;">
                <div style="background:white; padding:30px; border-radius:12px; box-shadow:0 4px 12px rgba(0,0,0,0.1); width:320px;">
                    <h2 style="text-align:center; color:#2c3e50; margin-bottom:20px;">🔐 Admin Login</h2>
                    <p style="color:red; text-align:center; margin-bottom:15px;">❌ Wrong username or password!</p>
                    <form method="POST">
                        <div style="margin-bottom:15px;">
                            <label style="display:block; margin-bottom:5px; font-weight:bold;">Username</label>
                            <input type="text" name="username" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                        </div>
                        <div style="margin-bottom:20px;">
                            <label style="display:block; margin-bottom:5px; font-weight:bold;">Password</label>
                            <input type="password" name="password" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                        </div>
                        <button type="submit" style="width:100%; padding:12px; background:#28a745; color:white; border:none; border-radius:6px; font-size:16px; cursor:pointer;">Login</button>
                    </form>
                </div>
            </body></html>
        """)
    return render_template_string("""
        <html><body style="background:#f0f4f8; font-family:Arial; display:flex; justify-content:center; align-items:center; height:100vh; margin:0;">
            <div style="background:white; padding:30px; border-radius:12px; box-shadow:0 4px 12px rgba(0,0,0,0.1); width:320px;">
                <h2 style="text-align:center; color:#2c3e50; margin-bottom:20px;">🔐 Admin Login</h2>
                <form method="POST">
                    <div style="margin-bottom:15px;">
                        <label style="display:block; margin-bottom:5px; font-weight:bold;">Username</label>
                        <input type="text" name="username" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                    </div>
                    <div style="margin-bottom:20px;">
                        <label style="display:block; margin-bottom:5px; font-weight:bold;">Password</label>
                        <input type="password" name="password" required style="width:100%; padding:10px; border:1px solid #ddd; border-radius:6px; box-sizing:border-box;">
                    </div>
                    <button type="submit" style="width:100%; padding:12px; background:#28a745; color:white; border:none; border-radius:6px; font-size:16px; cursor:pointer;">Login</button>
                </form>
            </div>
        </body></html>
    """)

@app.route('/logout')
def logout():
    resp = make_response('<script>window.location="/login";</script>')
    resp.set_cookie('logged_in', '', expires=0)
    return resp

# ===================== MAIN DASHBOARD =====================
@app.route('/')
def index():
    if not is_logged_in():
        return '<script>window.location="/login";</script>'
    return render_template_string("""
<!DOCTYPE html>
<html>
<head>
    <title>📚 Library Attendance System</title>
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <style>
        *{box-sizing:border-box; margin:0; padding:0; font-family:'Segoe UI', Arial, sans-serif;}
        body{background:#f0f4f8; padding:20px;}
        .container{max-width:1200px; margin:0 auto;}
        h1{text-align:center; color:#2c3e50; margin-bottom:25px;}
        .logout-btn{position:absolute; top:20px; right:20px; background:#dc3545; color:white; border:none; padding:10px 18px; border-radius:6px; cursor:pointer; font-size:15px;}
        .tabs{display:flex; flex-wrap:wrap; gap:10px; margin-bottom:25px;}
        .tab{padding:12px 20px; background:white; border:none; border-radius:8px; cursor:pointer; font-size:15px; font-weight:600; color:#555; transition:all 0.2s;}
        .tab:hover{background:#e9ecef;}
        .tab.active{background:#28a745; color:white;}
        .tab-content{background:white; padding:25px; border-radius:12px; box-shadow:0 2px 10px rgba(0,0,0,0.08); display:none;}
        .tab-content.active{display:block;}
        h2{color:#2c3e50; margin-bottom:20px; border-bottom:2px solid #eee; padding-bottom:10px;}
        input, select{width:100%; padding:12px; margin:8px 0 15px 0; border:1px solid #ddd; border-radius:6px; font-size:15px;}
        button{padding:12px 24px; background:#28a745; color:white; border:none; border-radius:6px; font-size:15px; cursor:pointer; margin:5px 5px 5px 0;}
        button:hover{background:#218838;}
        .btn-red{background:#dc3545;}
        .btn-red:hover{background:#c82333;}
        .btn-blue{background:#007bff;}
        .btn-blue:hover{background:#0056b3;}
        table{width:100%; border-collapse:collapse; margin-top:15px;}
        th, td{padding:12px; text-align:left; border-bottom:1px solid #eee;}
        th{background:#f8f9fa; font-weight:bold; color:#2c3e50;}
        .scan-box{text-align:center; padding:30px; background:#f8f9fa; border-radius:10px; margin-bottom:20px;}
        .scan-input{font-size:22px; text-align:center; padding:15px;}
        .success{color:green; font-weight:bold; margin-top:15px;}
        .error{color:red; font-weight:bold; margin-top:15px;}
        .barcode-img{max-width:250px; margin:15px 0;}
        .edit-form{background:#f8f9fa; padding:20px; border-radius:8px; margin-top:20px; display:none;}
    </style>
</head>
<body>
    <div class="container">
        <button class="logout-btn" onclick="window.location='/logout'">🚪 Logout</button>
        <h1>📚 Library Attendance System</h1>

        <div class="tabs">
            <button class="tab active" onclick="showTab('scan')">📱 Scan / Attendance</button>
            <button class="tab" onclick="showTab('register')">👤 Register Student</button>
            <button class="tab" onclick="showTab('students')">📋 All Students</button>
            <button class="tab" onclick="showTab('records')">📊 Attendance Records</button>
        </div>

        <!-- SCAN TAB -->
        <div id="scan" class="tab-content active">
            <h2>📱 Scan Student Barcode</h2>
            <div class="scan-box">
                <input type="text" id="scanInput" class="scan-input" placeholder="Scan barcode or type student number..." autofocus>
                <div id="scanResult"></div>
            </div>
        </div>

        <!-- REGISTER TAB -->
        <div id="register" class="tab-content">
            <h2>👤 Register New Student</h2>
            <form id="registerForm">
                <label><b>Full Name:</b></label>
                <input type="text" name="full_name" required placeholder="Juan Dela Cruz">
                
                <label><b>Department:</b></label>
                <select name="department" required>
                    <option value="">-- Select Department --</option>
                    {% for d in depts %}<option value="{{d}}">{{d}}</option>{% endfor %}
                </select>
                
                <label><b>Year Level:</b></label>
                <select name="year_level" required>
                    <option value="">-- Select Year --</option>
                    {% for y in years %}<option value="{{y}}">{{y}}</option>{% endfor %}
                </select>
                
                <label><b>Student Number:</b></label>
                <input type="text" name="student_number" required placeholder="e.g. 2024-0001">
                
                <label><b>Contact Number:</b></label>
                <input type="text" name="contact_number" placeholder="09123456789">
                
                <button type="submit">✅ Register & Generate Barcode</button>
            </form>
            <div id="registerResult"></div>
        </div>

        <!-- STUDENTS TAB -->
        <div id="students" class="tab-content">
            <h2>📋 Registered Students</h2>
            <button class="btn-blue" onclick="loadStudents()">🔄 Refresh List</button>
            <div id="studentsList"></div>
        </div>

        <!-- RECORDS TAB -->
        <div id="records" class="tab-content">
            <h2>📊 Attendance Records — Today</h2>
            <button class="btn-blue" onclick="loadRecords()">🔄 Refresh Records</button>
            <button class="btn-blue" onclick="exportRecords()">📄 Export to Word</button>
            <div id="recordsList"></div>
        </div>
    </div>

<script>
let currentEditId = null;

function showTab(tabName){
    document.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));
    document.querySelectorAll('.tab-content').forEach(c=>c.classList.remove('active'));
    event.target.classList.add('active');
    document.getElementById(tabName).classList.add('active');
    if(tabName==='students') loadStudents();
    if(tabName==='records') loadRecords();
    if(tabName==='scan') setTimeout(()=>document.getElementById('scanInput').focus(),100);
}

// SCAN / ATTENDANCE
document.getElementById('scanInput').addEventListener('keypress', function(e){
    if(e.key==='Enter') submitScan();
});

function submitScan(){
    const code = document.getElementById('scanInput').value.trim();
    if(!code) return;
    fetch('/scan', {
        method:'POST',
        headers:{'Content-Type':'application/x-www-form-urlencoded'},
        body:'student_number='+encodeURIComponent(code)
    }).then(r=>r.json()).then(d=>{
        const res = document.getElementById('scanResult');
        if(d.success){
            res.innerHTML = <div class="success">✅ ${d.message}</div>;
        }else{
            res.innerHTML = <div class="error">❌ ${d.message}</div>;
        }
        document.getElementById('scanInput').value='';
        setTimeout(()=>document.getElementById('scanInput').focus(),100);
    });
}

// REGISTER
document.getElementById('registerForm').addEventListener('submit', function(e){
    e.preventDefault();
    const fd = new FormData(this);
    fetch('/register', {method:'POST', body:fd})
    .then(r=>r.json()).then(d=>{
        const res = document.getElementById('registerResult');
        if(d.success){
            res.innerHTML = `<div class="success"><h3>✅ Registered Successfully!</h3>
                <p><b>Name:</b> ${d.student.full_name}</p>
                <p><b>Student No:</b> ${d.student.student_number}</p>
                <p><b>Barcode:</b></p>
                <img src="data:image/png;base64,${d.barcode}" class="barcode-img"><br>
                <button onclick="window.open('/print-barcode/${d.student.id}')">🖨️ Print Barcode</button>
                </div>`;
            this.reset();
        }else{
            res.innerHTML = <div class="error">❌ ${d.message}</div>;
        }
    });
});

// LOAD STUDENTS
function loadStudents(){
    fetch('/students').then(r=>r.text()).then(h=>{
        document.getElementById('studentsList').innerHTML = h;
    });
}

// LOAD RECORDS
function loadRecords(){
    fetch('/records').then(r=>r.text()).then(h=>{
        document.getElementById('recordsList').innerHTML = h;
    });
}

// EDIT STUDENT
function showEditForm(id, name, dept, year, num, contact){
    currentEditId = id;
    document.getElementById('edit_id').value = id;
    document.getElementById('edit_full_name').value = name;
    document.getElementById('edit_department').value = dept;
    document.getElementById('edit_year_level').value = year;
    document.getElementById('edit_student_number').value = num;
    document.getElementById('edit_contact_number').value = contact;
    document.getElementById('editForm').style.display = 'block';
}

function hideEditForm(){
    document.getElementById('editForm').style.display = 'none';
    currentEditId = null;
}

// EXPORT
function exportRecords(){
    window.location.href = '/export-word';
}
</script>

<!-- EDIT FORM -->
<div id="editForm" class="edit-form" style="display:none;">
    <h3>✏️ Edit Student Info</h3>
    <form id="editStudentForm">
        <input type="hidden" id="edit_id">
        <label><b>Full Name:</b></label>
        <input type="text" id="edit_full_name" required>
        <label><b>Department:</b></label>
        <select id="edit_department" required>
            <option value="CT">CT</option>
            <option value="FBT">FBT</option>
            <option value="BSED">BSED</option>
            <option value="BEED">BEED</option>
            <option value="BSFI">BSFI</option>
            <option value="BSBA">BSBA</option>
            <option value="EMPLOYEE">EMPLOYEE</option>
        </select>
        <label><b>Year Level:</b></label>
        <select id="edit_year_level" required>
            <option value="1st Year">1st Year</option>
            <option value="2nd Year">2nd Year</option>
            <option value="3rd Year">3rd Year</option>
            <option value="4th Year">4th Year</option>
            <option value="5th Year">5th Year</option>
            <option value="N/A">N/A</option>
        </select>
        <label><b>Student Number:</b></label>
        <input type="text" id="edit_student_number" required>
        <label><b>Contact Number:</b></label>
        <input type="text" id="edit_contact_number">
        <button type="button" onclick="saveEdit()">💾 Save Changes</button>
        <button type="button" class="btn-red" onclick="hideEditForm()">❌ Cancel</button>
    </form>
</div>

<script>
function saveEdit(){
    const id = document.getElementById('edit_id').value;
    const data = {
        id: id,
        full_name: document.getElementById('edit_full_name').value,
        department: document.getElementById('edit_department').value,
        year_level: document.getElementById('edit_year_level').value,
        student_number: document.getElementById('edit_student_number').value,
        contact_number: document.getElementById('edit_contact_number').value
    };
    fetch('/update-student', {
        method:'POST',
        headers:{'Content-Type':'application/json'},
        body: JSON.stringify(data)
    }).then(r=>r.json()).then(d=>{
        if(d.success){
            alert('✅ Updated successfully!');
            hideEditForm();
            loadStudents();
        }else{
            alert('❌ Error: ' + d.message);
        }
    });
}
</script>

</body>
</html>
    """, depts=DEPARTMENTS, years=YEAR_LEVELS)

# ===================== SCAN ENDPOINT =====================
@app.route('/scan', methods=['POST'])
def scan():
    if not is_logged_in():
        return jsonify({"success": False, "message": "Please login first"})
    
    student_number = request.form.get('student_number', '').strip().upper()
    now = datetime.datetime.now()
    today = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%I:%M %p")  # Oras lang, walang date
    
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    
    # Hanapin ang student
    c.execute("SELECT id, full_name FROM users WHERE UPPER(student_number) = %s", (student_number,))
    user = c.fetchone()
    if not user:
        conn.close()
        return jsonify({"success": False, "message": f"Student not found: {student_number}"})
    
    user_id, full_name = user
    
    # Check kung may active IN na walang OUT
    c.execute("SELECT id FROM attendance WHERE user_id = %s AND scan_date = %s AND time_out IS NULL", (user_id, today))
    active = c.fetchone()
    
    if active:
        # TIME OUT
        c.execute("UPDATE attendance SET time_out = %s WHERE id = %s", (time_str, active[0]))
        msg = f"⏰ TIME OUT — {full_name} — {time_str}"
    else:
        # TIME IN
        c.execute("INSERT INTO attendance (user_id, time_in, scan_date) VALUES (%s, %s, %s)", (user_id, time_str, today))
        msg = f"✅ TIME IN — {full_name} — {time_str}"
    
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": msg})

# ===================== REGISTER ENDPOINT =====================
@app.route('/register', methods=['POST'])
def register():
    if not is_logged_in():
        return jsonify({"success": False, "message": "Please login first"})
    
    full_name = request.form.get('full_name', '').strip()
    department = request.form.get('department', '').strip()
    year_level = request.form.get('year_level', '').strip()
    student_number = request.form.get('student_number', '').strip().upper()
    contact_number = request.form.get('contact_number', '').strip()
    registered_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    if not all([full_name, department, year_level, student_number]):
        return jsonify({"success": False, "message": "Fill all required fields"})
    
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    try:
        c.execute("""INSERT INTO users 
            (full_name, department, contact_number, year_level, student_number, registered_at)
            VALUES (%s, %s, %s, %s, %s, %s) RETURNING id""",
            (full_name, department, contact_number, year_level, student_number, registered_at))
        user_id = c.fetchone()[0]
        conn.commit()
        
        barcode_b64 = generate_barcode_img(student_number)
        
        return jsonify({
            "success": True,
            "student": {
                "id": user_id,
                "full_name": full_name,
                "student_number": student_number
            },
            "barcode": barcode_b64
        })
    except psycopg2.IntegrityError:
        conn.rollback()
        return jsonify({"success": False, "message": "Student Number already exists!"})
    finally:
        conn.close()

# ===================== STUDENTS LIST =====================
@app.route('/students')
def students_list():
    if not is_logged_in():
        return "Unauthorized"
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("SELECT id, full_name, department, year_level, student_number, contact_number FROM users ORDER BY full_name")
    students = c.fetchall()
    conn.close()
    
    html = "<table><tr><th>Name</th><th>Dept</th><th>Year</th><th>Student No.</th><th>Contact</th><th>Action</th></tr>"
    for s in students:
        html += f"""<tr>
            <td>{s[1]}</td><td>{s[2]}</td><td>{s[3]}</td><td>{s[4]}</td><td>{s[5] or '-'}</td>
            <td><button onclick="showEditForm({s[0]}, '{s[1]}', '{s[2]}', '{s[3]}', '{s[4]}', '{s[5] or ''}')">✏️ Edit</button>
            <a href="/print-barcode/{s[0]}" target="_blank"><button>🖨️ Barcode</button></a></td>
        </tr>"""
    html += "</table>"
    return html

# ===================== UPDATE STUDENT =====================
@app.route('/update-student', methods=['POST'])
def update_student():
    if not is_logged_in():
        return jsonify({"success": False})
    data = request.get_json()
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    try:
        c.execute("""UPDATE users SET full_name=%s, department=%s, year_level=%s, 
            student_number=%s, contact_number=%s WHERE id=%s""",
            (data['full_name'], data['department'], data['year_level'], 
             data['student_number'].upper(), data['contact_number'], data['id']))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": str(e)})
    finally:
        conn.close()

# ===================== ATTENDANCE RECORDS =====================
@app.route('/records')
def records():
    if not is_logged_in():
        return "Unauthorized"
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""SELECT u.full_name, u.department, a.time_in, a.time_out 
        FROM attendance a JOIN users u ON a.user_id = u.id 
        WHERE a.scan_date = %s ORDER BY a.id DESC""", (today,))
    recs = c.fetchall()
    conn.close()
    
    html = f"<h3>Attendance for {today}</h3><table><tr><th>Name</th><th>Dept</th><th>Time In</th><th>Time Out</th></tr>"
    for r in recs:
        html += f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[3] or '-'}</td></tr>"
    html += "</table>"
    return html

# ===================== EXPORT TO WORD =====================
@app.route('/export-word')
def export_word():
    if not is_logged_in():
        return "Unauthorized"
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("""SELECT u.full_name, u.department, a.time_in, a.time_out 
        FROM attendance a JOIN users u ON a.user_id = u.id 
        WHERE a.scan_date = %s ORDER BY a.id DESC""", (today,))
    recs = c.fetchall()
    conn.close()
    
    doc = Document()
    doc.add_heading(f'Library Attendance — {today}', 0)
    doc.add_paragraph(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Name'
    hdr[1].text = 'Department'
    hdr[2].text = 'Time In'
    hdr[3].text = 'Time Out'
    
    for r in recs:
        row = table.add_row().cells
        row[0].text = r[0]
        row[1].text = r[1]
        row[2].text = r[2]
        row[3].text = r[3] or '-'
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    resp = make_response(buffer.read())
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    resp.headers['Content-Disposition'] = f'attachment; filename=attendance_{today}.docx'
    return resp

# ===================== PRINT BARCODE =====================
@app.route('/print-barcode/<int:user_id>')
def print_barcode(user_id):
    if not is_logged_in():
        return "Unauthorized"
    conn = psycopg2.connect(DB_URI)
    c = conn.cursor()
    c.execute("SELECT full_name, student_number FROM users WHERE id = %s", (user_id,))
    user = c.fetchone()
    conn.close()
    if not user:
        return "User not found"
    
    barcode_b64 = generate_barcode_img(user[1])
    
    return f"""
    <html>
    <head><title>Barcode — {user[1]}</title>
    <style>
        body {{ text-align: center; padding: 40px; font-family: Arial; }}
        .barcode {{ max-width: 250px; margin: 20px auto; }}
        .info {{ font-size: 18px; margin: 10px 0; }}
        @media print {{ button {{ display: none; }} }}
    </style>
    </head>
    <body>
        <h2>{user[0]}</h2>
        <div class="info">Student No: {user[1]}</div>
        <img src="data:image/png;base64,{barcode_b64}" class="barcode">
        <br><br>
        <button onclick="window.print()" style="padding:12px 24px; font-size:16px; cursor:pointer;">🖨️ Print</button>
    </body>
    </html>
    """

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
